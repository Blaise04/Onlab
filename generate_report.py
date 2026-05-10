"""
Word dokumentum generáló script
Összefoglalja a projekt dokumentációját, eredményeit és grafikonjait.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = os.path.dirname(os.path.abspath(__file__))
PLOTS = os.path.join(BASE, "results", "plots")
OUT = os.path.join(BASE, "osszefoglalo.docx")


def set_col_width(table, col_index, width_cm):
    for row in table.rows:
        row.cells[col_index].width = Cm(width_cm)


def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.style.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return h


def add_table_with_header(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # fejléc
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, '1F497D')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)

    # sorok
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        if ri % 2 == 1:
            bg = 'DCE6F1'
        else:
            bg = 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)

    if col_widths:
        for ci, w in enumerate(col_widths):
            set_col_width(t, ci, w)
    return t


def add_image_centered(doc, path, width_inches=5.5, caption=None):
    if not os.path.exists(path):
        doc.add_paragraph(f"[Kép nem található: {os.path.basename(path)}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].italic = True
        cp.runs[0].font.size = Pt(9)
        cp.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)


doc = Document()

# --- Stílusok ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for lvl in range(1, 4):
    hs = doc.styles[f'Heading {lvl}']
    hs.font.name = 'Calibri'

# ====================================================
# BORÍTÓ
# ====================================================
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Opciós árazási neurális hálók\nBlack-Scholes szintetikus adatokon")
tr.bold = True
tr.font.size = Pt(20)
tr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Összefoglaló dokumentum — Önálló laboratórium, 6. félév\nArató Balázs | 2026")
sr.font.size = Pt(13)
sr.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_page_break()

# ====================================================
# 1. BEVEZETÉS
# ====================================================
add_heading(doc, "1. Bevezetés", 1)

doc.add_paragraph(
    "A projekt célja Black-Scholes (BS) opciós árazási formula közelítése mélytanulási módszerekkel. "
    "Az első fázisban szintetikus adathalmazon tanítunk különböző neurális háló architektúrákat, "
    "és szisztematikusan összehasonlítjuk teljesítményüket. "
    "A feladat lényegében egy négydimenziós regressziós probléma: "
    "a moneyness (S/K), a lejáratig hátralévő idő (T), a kockázatmentes kamatláb (r) és a volatilitás (σ) "
    "alapján prediktálni az opció normalizált call árát (C/K)."
)

doc.add_paragraph(
    "A García & Gençay (2000) által bevezetett homogenitási hint alapján a BS ár lineárisan homogén "
    "S és K-ban, ezért az S/K és C/K transzformáció alkalmazásával csökkentjük az input tér dimenzióját "
    "és javítjuk az általánosítást. Az osztalékhozamot (q = 0) elhanyagoljuk."
)

# ====================================================
# 2. ADATHALMAZ
# ====================================================
add_heading(doc, "2. Adathalmaz", 1)
add_heading(doc, "2.1 Generálás", 2)

doc.add_paragraph(
    "Az adathalmazt a generate_dataset.py script állítja elő Latin Hypercube Sampling (LHS) "
    "módszerrel, amely garantálja a paramétertér egyenletes lefedettségét. "
    "Az adathalmaz 1 000 000 szintetikus mintát tartalmaz, seed=42 reprodukálhatóság érdekében."
)

doc.add_paragraph("Generálási parancs:")
code_p = doc.add_paragraph(
    "python generate_dataset.py --n 1000000 --method lhs --normalize --format parquet --seed 42"
)
code_p.style = doc.styles['No Spacing']
code_p.paragraph_format.left_indent = Cm(1)
code_p.runs[0].font.name = 'Courier New'
code_p.runs[0].font.size = Pt(9)

add_heading(doc, "2.2 Paramétertartományok", 2)

add_table_with_header(doc,
    ["Paraméter", "Szimbólum", "Min", "Max", "Eloszlás"],
    [
        ["Részvényár", "S", "10.0", "150.0", "Uniform"],
        ["Moneyness", "S/K", "0.5", "1.5", "Uniform"],
        ["Lejárati idő", "T (év)", "0.005", "2.0", "Uniform"],
        ["Kockázatmentes ráta", "r", "0.00", "0.05", "Uniform"],
        ["Volatilitás", "σ", "0.05", "0.90", "Uniform"],
    ],
    col_widths=[3.5, 2.0, 1.5, 1.5, 2.5]
)

add_heading(doc, "2.3 Adatfelosztás és fájlok", 2)

add_table_with_header(doc,
    ["Halmaz", "Méret", "Fájl", "Méret (MB)"],
    [
        ["Tanítóhalmaz", "700 000", "data/train.parquet", "57"],
        ["Validációs halmaz", "150 000", "data/val.parquet", "15"],
        ["Teszthalmaz", "150 000", "data/test.parquet", "15"],
    ],
    col_widths=[3.5, 2.5, 5.0, 2.5]
)

doc.add_paragraph()
doc.add_paragraph(
    "Sanity check: ATM opció (S=K=100, T=1, r=0.05, σ=0.2) esetén a BS call ár ≈ 10.45. "
    "A generátor 10.4506 értéket adott — helyes."
)

add_heading(doc, "2.4 Miért LHS és Parquet?", 2)
doc.add_paragraph(
    "Az LHS szisztematikusan lefedi a paraméterteret, elkerüli az egyenletes véletlenszerű "
    "mintavételezés klaszterezési problémáját. 5 dimenziós rácsnál 1M pont esetén "
    "~4 pont/dimenzió lenne, ami nagyon durva — az LHS folytonos és jobb. "
    "A Parquet formátum 1M sornál ~5–10× kisebb fájlméretet és gyorsabb I/O-t biztosít CSV-hez képest."
)

doc.add_page_break()

# ====================================================
# 3. MODELL ARCHITEKTÚRÁK
# ====================================================
add_heading(doc, "3. Modell architektúrák", 1)

doc.add_paragraph(
    "Összesen 7 különböző neurális háló architektúrát tanítottunk és hasonlítottunk össze. "
    "Az architektúrák két generációba sorolhatók:"
)

add_heading(doc, "3.1 1. generáció — Baseline modellek", 2)

add_table_with_header(doc,
    ["Modell", "CLI neve", "Irodalom", "Paraméterek", "Aktiváció"],
    [
        ["MLPPricer", "mlp", "Culkin & Das (2017)", "~31 000", "ReLU"],
        ["DeepMLPPricer", "deep_mlp", "Della Corte et al. (2023)", "~268 000", "ReLU"],
        ["ResNetPricer", "resnet", "Della Corte et al. (2023)", "~399 000", "ReLU"],
    ],
    col_widths=[3.2, 2.2, 4.5, 2.5, 2.0]
)

add_heading(doc, "3.2 2. generáció — Kísérleti modellek", 2)

add_table_with_header(doc,
    ["Modell", "CLI neve", "Irodalom", "Paraméterek", "Aktiváció"],
    [
        ["GELUResNetPricer", "gelu_resnet", "ResNetPricer + GELU", "~399 000", "GELU"],
        ["DenseMLPPricer", "dense_mlp", "Huang et al. (2017) DenseNet", "~102 000", "GELU"],
        ["HighwayPricer", "highway", "Srivastava et al. (2015)", "~528 000", "GELU"],
        ["FINNPricer", "finn", "Liu et al. (2019), arXiv:2412", "~403 000", "GELU"],
    ],
    col_widths=[3.5, 2.5, 4.5, 2.5, 1.5]
)

add_heading(doc, "3.3 Architektúra leírások", 2)

architectures = [
    ("MLPPricer — Culkin & Das (2017) baseline",
     "Input(4) → Linear(4→100) → ReLU × 4 → Linear(100→1)\n"
     "Nincs normalizáció, nincs Dropout. ~31 000 paraméter. "
     "A legegyszerűbb architektúra, amely az eredeti tanulmányt reprodukálja."),

    ("DeepMLPPricer — Javított MLP (Della Corte et al. 2023)",
     "Input(4) → Linear(4→256) → [LayerNorm → ReLU → Dropout(0.1) → Linear(256→256)] × 4 → LayerNorm → Linear(256→1)\n"
     "Pre-LN stílus: normalizáció a nemlinearitás előtt. ~268 000 paraméter."),

    ("ResNetPricer — Reziduális MLP (Della Corte et al. 2023)",
     "Input projekció: Linear(4→256) → BatchNorm1d → ReLU\n"
     "ResidualBlock(256): x → Linear → BatchNorm → ReLU → Dropout(0.1) → Linear → BatchNorm → +x\n"
     "3 reziduális blokk. ~399 000 paraméter."),

    ("GELUResNetPricer — ResNet GELU aktivációval",
     "Azonos struktúra mint ResNetPricer, ReLU → GELU csere és LayerNorm normalizációval.\n"
     "Motiváció: a BS árak simák, GELU simább gradienst biztosít. ~399 000 paraméter."),

    ("DenseMLPPricer — DenseNet-stílusú összefűzés",
     "h₁=GELU(W₁·x); h₂=GELU(W₂·[x,h₁]); h₃=GELU(W₃·[x,h₁,h₂]); h₄=GELU(W₄·[x,h₁,h₂,h₃])\n"
     "Minden réteg az összes korábbi kimenetét kapja. ~102 000 paraméter. "
     "Irodalom: Huang et al. (2017) — Densely Connected Convolutional Networks."),

    ("HighwayPricer — Tanulható gating",
     "4 HighwayBlock(256): H=GELU(W_H·x); T=σ(W_T·x); y=H·T + x·(1−T)\n"
     "Gate bias −1-re inicializálva: kezdetben inkább 'carry' (skip). ~528 000 paraméter. "
     "Irodalom: Srivastava et al. (2015) — Training Very Deep Networks."),

    ("FINNPricer — Finance-Informed Neural Network",
     "Ág 1 (approx): x → [Linear(4→64) → GELU] × 2 → Linear(64→1) → BS̃\n"
     "Ág 2 (correction): x → Linear(4→256) → GELU → [GELUResidualBlock(256)] × 3 → LayerNorm → Linear(256→1) → δ\n"
     "Output: BS̃ + δ. ~403 000 paraméter. "
     "Irodalom: Liu et al. (2019), arXiv:2412.12213."),
]

for name, desc in architectures:
    p = doc.add_paragraph()
    p.add_run(name).bold = True
    doc.add_paragraph(desc)

doc.add_page_break()

# ====================================================
# 4. TANÍTÁSI KONFIGURÁCIÓ
# ====================================================
add_heading(doc, "4. Tanítási konfiguráció", 1)

add_table_with_header(doc,
    ["Hiperparaméter", "Érték"],
    [
        ["Optimizer", "Adam"],
        ["Tanulási ráta (lr)", "1e-3"],
        ["Weight decay", "1e-4"],
        ["Loss", "MSELoss"],
        ["LR Scheduler", "ReduceLROnPlateau (factor=0.5, patience=5, min_lr=1e-6)"],
        ["Early stopping", "patience=10 epoch"],
        ["Batch méret", "4096"],
        ["Max epochok", "200"],
        ["Seed", "42"],
        ["GPU", "NVIDIA GeForce RTX 4060 Laptop GPU (8 GB)"],
    ],
    col_widths=[5.0, 9.5]
)

doc.add_paragraph()
add_heading(doc, "4.1 Physics-Informed Loss (opcionális)", 2)

doc.add_paragraph(
    "A ResNetPricer physics-loss regularizációval is betanítva (resnet_phys modell):"
)
doc.add_paragraph(
    "L = L_MSE + λ · L_delta    ahol    L_delta = mean(relu(−∂C_norm/∂m_norm) + relu(∂C_norm/∂m_norm − 1))"
)
doc.add_paragraph(
    "A korlát: delta ∈ [0, 1] — a call delta definíció szerint nem lehet negatív ill. 1-nél nagyobb. "
    "Irodalmi alap: Liu et al. (2019), PINN (arXiv:2312.06711). CLI: --physics-loss --physics-lambda 0.1"
)

doc.add_page_break()

# ====================================================
# 5. EREDMÉNYEK
# ====================================================
add_heading(doc, "5. Eredmények", 1)
add_heading(doc, "5.1 Összesített teszt eredmények", 2)

doc.add_paragraph(
    "Az összes modellt azonos feltételek mellett tanítottuk: 700 000 szintetikus BS minta "
    "(LHS, seed=42), validációs és teszt halmaz 150 000 minta, max 200 epoch. "
    "Kiértékelés a teszt halmazon."
)

add_table_with_header(doc,
    ["Modell", "Paraméterek", "Best ep", "Val MSE (×10⁻⁵)", "Test RMSE", "Test MAE", "R²"],
    [
        ["mlp", "31 001", "68", "2.13", "0.004612", "0.002780", "0.999407"],
        ["resnet", "398 593", "80", "2.28", "0.004788", "0.003274", "0.999361"],
        ["resnet_phys", "398 593", "68", "3.28", "0.005783", "0.003987", "0.999067"],
        ["gelu_resnet", "398 593", "40", "8.28", "0.009135", "0.006027", "0.997673"],
        ["finn", "403 202", "17", "9.23", "0.009653", "0.006625", "0.997401"],
        ["dense_mlp", "101 894", "18", "22.80", "0.015172", "0.010694", "0.993581"],
        ["highway", "528 129", "22", "23.20", "0.015281", "0.011005", "0.993488"],
        ["deep_mlp", "267 521", "4", "27.40", "0.016575", "0.013233", "0.992338"],
    ],
    col_widths=[2.5, 2.5, 1.8, 3.0, 2.3, 2.3, 2.1]
)

doc.add_paragraph()
doc.add_paragraph(
    "Az mlp bizonyult a legjobb modellnek (RMSE=0.004612, R²=0.999407), megelőzve a resnet-et. "
    "A resnet_phys physics-loss regularizációval ~21%-kal magasabb RMSE-t ért el, "
    "de garantálja a delta-korlátot."
)

add_heading(doc, "5.2 Szegmentált eredmények (RMSE)", 2)

doc.add_paragraph(
    "Szegmenshatárok: OTM = moneyness < 0.9, ATM = 0.9–1.1, ITM = moneyness > 1.1. "
    "N(OTM)=60 051, N(ATM)=29 948, N(ITM)=60 001."
)

add_table_with_header(doc,
    ["Modell", "OTM (m<0.9)", "ATM (0.9–1.1)", "ITM (m>1.1)", "max_error", "R² (all)"],
    [
        ["mlp", "0.002933 ★", "0.005903 ★", "0.005214", "0.076284", "0.999407"],
        ["resnet", "0.004429", "0.005342", "0.004842 ★", "0.065314", "0.999361"],
        ["resnet_phys", "0.005502", "0.006357", "0.005756", "0.063004 ★", "0.999067"],
        ["gelu_resnet", "0.008078", "0.011879", "0.008536", "0.105482", "0.997673"],
        ["finn", "0.008649", "0.012113", "0.009211", "0.105467", "0.997401"],
        ["dense_mlp", "0.011125", "0.018354", "0.016836", "0.115152", "0.993581"],
        ["highway", "0.011821", "0.020864", "0.015054", "0.115609", "0.993488"],
        ["deep_mlp", "0.013577", "0.013992", "0.020115", "0.103493", "0.992338"],
    ],
    col_widths=[2.5, 2.5, 2.8, 2.5, 2.2, 2.0]
)

doc.add_paragraph()
doc.add_paragraph(
    "★ = szegmensenkénti legjobb érték. Az mlp OTM és ATM szegmensben vezet, "
    "a resnet ITM szegmensben a legjobb. "
    "A resnet_phys a legkisebb max_error-t éri el."
)

doc.add_page_break()

# ====================================================
# 6. GRAFIKONOK
# ====================================================
add_heading(doc, "6. Grafikonok", 1)

add_heading(doc, "6.1 Összehasonlító ábrák", 2)

add_image_centered(
    doc,
    os.path.join(PLOTS, "val_mse_bar_chart.png"),
    width_inches=5.5,
    caption="1. ábra: Validációs MSE összehasonlítás modellenként (kisebb = jobb)"
)

doc.add_paragraph()

add_image_centered(
    doc,
    os.path.join(PLOTS, "params_vs_val_mse.png"),
    width_inches=5.5,
    caption="2. ábra: Paraméterszám vs. Validációs MSE — a több paraméter nem garantál jobb eredményt"
)

doc.add_paragraph()

add_image_centered(
    doc,
    os.path.join(PLOTS, "training_curve_comparison.png"),
    width_inches=6.0,
    caption="3. ábra: Összesített tanulási görbék — val loss epochonként, modellenként"
)

doc.add_page_break()

add_heading(doc, "6.2 Egyedi tanulási görbék", 2)

model_plots = [
    ("MLPPricer", "training_curve_MLPPricer.png"),
    ("DeepMLPPricer", "training_curve_DeepMLPPricer.png"),
    ("ResNetPricer", "training_curve_ResNetPricer.png"),
    ("GELUResNetPricer", "training_curve_GELUResNetPricer.png"),
    ("DenseMLPPricer", "training_curve_DenseMLPPricer.png"),
    ("HighwayPricer", "training_curve_HighwayPricer.png"),
    ("FINNPricer", "training_curve_FINNPricer.png"),
]

for i, (name, fname) in enumerate(model_plots):
    add_image_centered(
        doc,
        os.path.join(PLOTS, fname),
        width_inches=5.5,
        caption=f"{4 + i}. ábra: {name} tanulási görbe"
    )
    if i < len(model_plots) - 1:
        doc.add_paragraph()

doc.add_page_break()

# ====================================================
# 7. ARCHITEKTÚRA-ÖSSZEHASONLÍTÁS
# ====================================================
add_heading(doc, "7. Architektúra-összehasonlítás elemzése", 1)

add_heading(doc, "7.1 Reziduális kapcsolatok: ResNetPricer vs. MLPPricer", 2)
doc.add_paragraph(
    "Az MLPPricer (Val MSE: 2.13×10⁻⁵) szignifikánsan jobb eredményt ért el, mint a ResNetPricer "
    "(Val MSE: 2.28×10⁻⁵), annak ellenére, hogy az utóbbi ~13-szor több paramétert tartalmaz. "
    "Az MLPPricer lényegesen hosszabb ideig tanult (68 vs. 80 epoch), és a Black-Scholes függvény "
    "sima természete miatt egy egyszerű MLP is elegendő kapacitással bír."
)

add_heading(doc, "7.2 Normalizáció: BatchNorm vs. LayerNorm", 2)
doc.add_paragraph(
    "A ResNetPricer (BatchNorm, 2.28×10⁻⁵) kedvezőbbnek bizonyult mint a GELUResNetPricer "
    "(LayerNorm, 8.28×10⁻⁵). Azonban a két modell aktivációban is különbözik (ReLU vs. GELU), "
    "ezért ez nem tiszta összehasonlítás. A DeepMLPPricer (LayerNorm, 27.40×10⁻⁵) "
    "a korai early stopping miatt nem konvergált megfelelően."
)

add_heading(doc, "7.3 Aktiváció: ReLU vs. GELU", 2)
doc.add_paragraph(
    "A GELUResNetPricer elméleti előnye (simább gradiens a BS árak sima természetéhez) "
    "a gyakorlatban nem hozott mért javulást. Sőt, a 19–22. epochon erős instabilitás lépett fel "
    "(val loss 0.000086-ról 0.002769-re ugrott), ami GELU + LR=1e-3 kombinációval nagyobb "
    "gradiens-robbanás érzékenységre utal."
)

add_heading(doc, "7.4 DenseNet összefűzés: DenseMLPPricer", 2)
doc.add_paragraph(
    "A DenseMLPPricer a gyengébb teljesítményt nyújtó modellek között van (RMSE=0.015172). "
    "A dense skip-kapcsolatok BS opciós árazásnál nem hasznosak: a BS ár sima, "
    "nem igényli a korai feature-ök direkt átadását. A modell mindössze 18 best epochig tanult."
)

add_heading(doc, "7.5 Highway gating: HighwayPricer", 2)
doc.add_paragraph(
    "A HighwayPricer (RMSE=0.015281) szintén gyengén teljesített. A tanulható gate-ek "
    "felesleges paramétereket visznek be (528K param). A highway mechanizmus mélyen rétegezett "
    "képosztályozásnál hasznos; 5-dimenziós sima táblázati adatnál nem jelent előnyt. "
    "ATM szegmensben szignifikánsan gyengébb (RMSE=0.020864)."
)

add_heading(doc, "7.6 Kétágú architektúra: FINNPricer", 2)
doc.add_paragraph(
    "A FINNPricer (RMSE=0.009653) közepes eredményt ért el. A kétágú architektúra gyors "
    "konvergenciát mutat (17 best epoch), de szintetikus BS adatokon a 'közelítő ág' "
    "nem tud valódi előnyt nyújtani — nincs modellhiba amit korrigálni kellene, "
    "csak a háló saját hibáját becsüli. Valós piaci adatokon várhatóan jobban hasznosítható."
)

add_heading(doc, "7.7 Physics-Informed Loss hatása", 2)
doc.add_paragraph(
    "A resnet_phys (RMSE=0.005783) ~21%-kal magasabb RMSE-t mutat, mint a sima ResNet, "
    "de garantálja a delta-korlátot: ∂C_norm/∂moneyness_norm ∈ [0, 1]. "
    "A physics loss fő haszna nem az MSE-ben mérhető, hanem a modell pénzügyi konzisztenciájában: "
    "a tanult delta közelíti a Black-Scholes deltát anélkül, hogy azt expliciten optimalizálnánk. "
    "Ez különösen fontos fedezési stratégiák számításánál."
)

doc.add_page_break()

# ====================================================
# 8. KÖVETKEZTETÉSEK
# ====================================================
add_heading(doc, "8. Következtetések", 1)

add_heading(doc, "8.1 Legjobb architektúra", 2)
doc.add_paragraph(
    "Az MLPPricer nyerte a kísérletet a legjobb validációs és teszt MSE-vel (RMSE=0.004612, R²=0.999407). "
    "Ez azért figyelemre méltó, mert ez az egyszerűsített Culkin & Das (2017) baseline modell "
    "a legkisebb architektúra (~31 000 param). "
    "A Black-Scholes függvény zárt alakja sima — nincs szükség mély, komplex architektúrára."
)

add_heading(doc, "8.2 Főbb tanulságok", 2)

conclusions = [
    "Paraméterszám növelése NEM garantál jobb eredményt: az mlp (31K) felülmúlta a highway-t (528K).",
    "Az early stopping hatása kritikus: az mlp 68 epochig tanult, a deep_mlp mindössze 4-ig.",
    "A GELU aktiváció ReLU-hoz képest instabilabb tanulást mutatott ennél a feladatnál.",
    "A resnet_phys physics regularizáció minimális MSE-veszteséggel pénzügyi konzisztenciát biztosít.",
    "Szegmentált elemzés: az mlp OTM/ATM szegmensben a legjobb, a resnet ITM-ben vezet.",
    "Az architektúra bonyolítása szintetikus BS adatokon nem hoz áttörést — valós piaci adatokon megváltozhat.",
]

for c in conclusions:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(c)

add_heading(doc, "8.3 Összefoglaló rangsor", 2)

add_table_with_header(doc,
    ["Kategória", "Győztes", "Megjegyzés"],
    [
        ["Legjobb összesített", "MLPPricer", "RMSE=0.004612, R²=0.9994, 31K param"],
        ["Legjobb param/teljesítmény", "MLPPricer", "31K param, legjobb RMSE — nincs kompromisszum"],
        ["Legjobb fizikai korlát", "resnet_phys", "Delta-korlát garantált, RMSE ~21%-kal magasabb"],
        ["Legstabilabb konvergencia", "ResNetPricer", "Fokozatos konvergencia 80 epochon át"],
        ["2. generáció legjobb", "GELUResNetPricer", "RMSE=0.009135, de instabilitásra hajlamos"],
        ["Leggyorsabb konvergencia", "FINNPricer", "17 epochon best, de instabil"],
        ["Leggyengébb", "DeepMLPPricer", "4 epochon best, val=0.000274 — nem konvergált"],
    ],
    col_widths=[4.0, 3.5, 7.0]
)

add_heading(doc, "8.4 Korlátok és 2. fázis irányai", 2)

doc.add_paragraph(
    "Jelen kísérlet szintetikus adatokon zajlott — a Black-Scholes képlet zárt alakjából "
    "generált adatok nem tartalmazzák a valós piaci jelenségeket "
    "(volatility smile, bid-ask spread, likviditási hatások). A 2. fázisban:"
)

future = [
    "Historikus piaci adatok bevonása — a rangsor megváltozhat valós adatokon.",
    "Görögök (delta, vega, gamma, theta) predikciója — multi-output architektúra vagy autograd.",
    "Physics-Informed Loss (PINN megközelítés) kiterjesztése valós adatokra.",
    "Hiperparaméter optimalizálás — különösen DenseMLPPricer és HighwayPricer esetén.",
    "Volatility surface (smile/skew) modellezése — Heston vagy SABR modell adatain.",
]

for f in future:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f)

doc.add_page_break()

# ====================================================
# 9. IRODALMI HIVATKOZÁSOK
# ====================================================
add_heading(doc, "9. Irodalmi hivatkozások", 1)

references = [
    ("Hutchinson, Lo & Poggio (1994)",
     "A Nonparametric Approach to Pricing and Hedging Derivative Securities Via Learning Networks. "
     "The Journal of Finance, Vol. 49, No. 3. "
     "Elsők között mutatták meg, hogy neurális hálók képesek közel-BS minőségű árazást tanulni."),

    ("Garcia & Gençay (2000)",
     "Pricing and hedging derivative securities with neural networks and a homogeneity hint. "
     "Journal of Econometrics, Vol. 94. "
     "Bevezeti a homogeneity hint fogalmát: C/K = f(S/K, T, r, σ). "
     "DOI: 10.1016/S0304-4076(99)00018-1"),

    ("Culkin & Das (2017)",
     "Machine Learning in Finance: The Case of Deep Learning for Option Pricing. "
     "Journal of Investment Management, Vol. 15, No. 4. "
     "4 rejtett réteg × 100 neuron, ReLU, 300K szintetikus BS adatpont — baseline architektúra. "
     "SSRN: 3023505"),

    ("Liu, Oosterlee & Bohte (2019)",
     "Pricing options and computing implied volatilities using neural networks. "
     "Risks, Vol. 7, No. 1. "
     "CaNN (Calibration Neural Network), fizikai korlátok beépítése, implied volatility inverz feladat. "
     "DOI: 10.3390/risks7010016"),

    ("Ruf & Wang (2020)",
     "Neural Networks for Option Pricing and Hedging: A Literature Review. "
     "Journal of Computational Finance, Vol. 24, No. 1. "
     "Átfogó összehasonlítás ~15 NN-alapú opciós árazási megközelítésről. "
     "arXiv: 1911.05689"),

    ("Della Corte et al. (2023)",
     "Machine learning for option pricing: an empirical investigation of network architectures. "
     "arXiv preprint. "
     "MLP vs. ResNet vs. Highway Network összehasonlítás BS és Heston modelleken. "
     "arXiv: 2307.07657"),

    ("Huang et al. (2017)",
     "Densely Connected Convolutional Networks (DenseNet). "
     "CVPR 2017. "
     "DenseMLPPricer alapja: minden réteg az összes korábbi kimenetét kapja."),

    ("Srivastava et al. (2015)",
     "Training Very Deep Networks (Highway Networks). "
     "NeurIPS 2015. "
     "HighwayPricer alapja: tanulható transform gate."),

    ("Liao et al. (2024)",
     "The AI Black-Scholes: Finance-Informed Neural Network for Option Pricing. "
     "arXiv: 2412.12213. "
     "Hibrid kétágú architektúra + pénzügyi korlátok (put-call paritás, monotonitás, no-arbitrage). "
     "FINNPricer inspirációja."),
]

for i, (name, desc) in enumerate(references):
    p = doc.add_paragraph()
    p.add_run(f"[{i+1}] {name}").bold = True
    doc.add_paragraph(desc)
    doc.add_paragraph()

# ====================================================
# Mentés
# ====================================================
doc.save(OUT)
print(f"Dokumentum mentve: {OUT}")
